import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="PEDATI Master Planner", layout="wide")
st.title("🎓 PEDATI SMART LESSON PLANNER")

# --- MAIN PAGE CONFIGURATION & USER API KEY BAR (AT THE VERY TOP) ---
user_api_key = st.text_input(
    "🔑 ENTER YOUR GEMINI API KEY:", 
    type="password", 
    help="Get your API key from Google AI Studio using your Gmail account."
)

# Helper function to dynamically check and load models based on the user's key
def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if "flash" in m.name:
                    return m.name
        return "models/gemini-1.5-flash"
    except Exception as e:
        st.error(f"INVALID API KEY OR CONNECTION ERROR: {str(e)}")
        return None

# Process model assignment if the key is provided
selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.info(f"SYSTEM CONNECTED VIA YOUR API KEY. ACTIVE MODEL: {selected_model_name.upper()}")
else:
    st.warning("⚠️ PLEASE ENTER YOUR PERSONAL GEMINI API KEY ABOVE TO START.")


# --- 2. AI GENERATION ENGINE ---
def generate_pedati_plan(topic, syllabus, extra_context, api_key, model_name):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    Topic: {topic}. Syllabus Code: {syllabus}. Context: {extra_context}.
    Generate a lesson plan in English. 
    NO Malay terms.
    
    CRITICAL TEXT FORMATTING RULES:
    1. DO NOT use double asterisks (**) anywhere in the output.
    2. Ensure every single section title and stage marker is in full CAPITAL LETTERS.
    3. Use numbers for lists where appropriate.
    
    Use these exact stage names:
    P [Prior Knowledge], E [Engage], D [Develop], A [Apply], T [Test], I [Improve].

    Structure with these exact markers for boxing:
    SECTION: LESSON OBJECTIVES
    [4 points]
    SECTION: LESSON OUTCOMES
    [4 points]
    SECTION: SUCCESS CRITERIA
    [4 points]
    SECTION: PREREQUISITE
    [1 point]
    SECTION: KEYWORDS
    [6 items]
    SECTION: HOTS
    [any 4 main domains in the Bloom's taxonomy]
    SECTION: DIGITAL CITIZENSHIP
    [4 points on the use of online resources like youtube channel or canva application or use of chromebooks or use of digital devices]

    SECTION: PEDATI STAGES
    STAGE: P [PRIOR KNOWLEDGE] | CB: [Activity] | SB: [Activity]
    STAGE: E [ENGAGE] | CB: [Activity] | SB: [Activity]
    STAGE: D [DEVELOP] | CB: [Activity] | SB: [Activity]
    STAGE: A [APPLY] | CB: [Activity] | SB: [Activity]
    STAGE: T [TEST] | CB: [Activity] | SB: [Activity]
    STAGE: I [IMPROVE] | CB: [Activity] | SB: [Activity]
    """
    try:
        response = model.generate_content(prompt)
        return response.text.replace("**", "")
    except Exception as e:
        return f"SYSTEM ERROR: {str(e)}"


# --- Helper Function for Word XML Native Page Numbers ---
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


# --- 3. WORD DOCUMENT EXPORT ENGINE (RECONFIGURED) ---
def create_word_export(topic, syllabus, text):
    doc = Document()
    
    # 1. Page Size Setup (Letter dimensions: 8.5" x 11.5") & 0.5" Margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # 2. Add Native Page Numbering to Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = 2  # Right aligned
    footer_run = footer_p.add_run("Page ")
    footer_run.font.name = 'Arial Narrow'
    footer_run.font.size = Pt(10)
    add_page_number(footer_run)

    # 3. Global Dynamic Typography & Spacing Configurations (Base Content: 12pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(12)
    
    p_format = style.paragraph_format
    p_format.line_spacing = 1.0  # Tight single spacing
    p_format.space_after = Pt(0)   # Clean compression
    p_format.space_before = Pt(0)

    # Document Header Title - 14PT, BOLD & CAPITAL LETTERS
    title_p = doc.add_paragraph()
    title_text = f'LESSON PLAN: {topic.upper()} ({syllabus.upper()})'
    run_title = title_p.add_run(title_text)
    run_title.bold = True
    run_title.font.size = Pt(14)
    title_p.paragraph_format.space_after = Pt(6)

    # Admin Header Table (6-field layout)
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["WEEK NO:", "DATE:"], ["NO. OF STUDENTS:", "DAY:"], ["VENUE / LAB NO:", "DURATION (MINS):"]]
    for r in range(3):
        for col_idx, text_lbl in [(0, labels[r][0]), (2, labels[r][1])]:
            cell_p = admin_table.cell(r, col_idx).paragraphs[0]
            cell_p.paragraph_format.line_spacing = 1.0
            cell_p.paragraph_format.space_after = Pt(0)
            run = cell_p.add_run(text_lbl)
            run.bold = True
            run.font.size = Pt(12)
    
    # Spacer paragraph
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    # Resources Table Header - 14PT & CAPITAL LETTERS
    p_res = doc.add_paragraph()
    run_res_hdr = p_res.add_run("RESOURCES & MATERIALS")
    run_res_hdr.bold = True
    run_res_hdr.font.size = Pt(14)
    p_res.paragraph_format.space_after = Pt(4)
    
    res_table = doc.add_table(rows=1, cols=1)
    res_table.style = 'Table Grid'
    res_cell_p = res_table.cell(0, 0).paragraphs[0]
    res_cell_p.paragraph_format.line_spacing = 1.0
    res_cell_p.paragraph_format.space_after = Pt(0)
    res_run = res_cell_p.add_run("Smart board, Chromebook, Writing table, Projector, Screen share with laptop")
    res_run.font.size = Pt(12)
    
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_after = Pt(6)

    # Content Parsing & Table Boxing
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): 
            continue
            
        lines = section.strip().split('\n')
        title = lines[0].strip().upper().replace("**", "")
        content_lines = lines[1:]

        # Section Heading - 14PT, BOLD & CAPITAL LETTERS
        p_sec = doc.add_paragraph()
        run_sec_title = p_sec.add_run(title)
        run_sec_title.bold = True
        run_sec_title.font.size = Pt(14)
        p_sec.paragraph_format.space_after = Pt(4)

        if "|" in section and "PEDATI" in title:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            
            # Table Header Content - 12PT Bold
            for idx, h_text in enumerate(['STAGE (PEDATI)', ' ACTIVITY 1 ', ' ACTIVITY 2 ']):
                hdr_p = hdr[idx].paragraphs[0]
                hdr_p.paragraph_format.line_spacing = 1.0
                hdr_p.paragraph_format.space_after = Pt(0)
                hrun = hdr_p.add_run(h_text)
                hrun.bold = True
                hrun.font.size = Pt(12)

            for line in content_lines:
                cleaned_line = line.replace("**", "").strip()
                if "|" in cleaned_line:
                    p = cleaned_line.split("|")
                    if len(p) >= 3:
                        row_cells = table.add_row().cells
                        
                        col_0 = p[0].split(":")[-1].strip().upper() if ":" in p[0] else p[0].strip().upper()
                        col_1 = p[1].split(":")[-1].strip() if ":" in p[1] else p[1].strip()
                        col_2 = p[2].split(":")[-1].strip() if ":" in p[2] else p[2].strip()
                        
                        for c_idx, c_text in enumerate([col_0, col_1, col_2]):
                            cell_p = row_cells[c_idx].paragraphs[0]
                            cell_p.paragraph_format.line_spacing = 1.0
                            cell_p.paragraph_format.space_after = Pt(0)
                            c_run = cell_p.add_run(c_text)
                            c_run.font.size = Pt(12)
                            if c_idx == 0:
                                c_run.bold = True
            
            s_par = doc.add_paragraph()
            s_par.paragraph_format.space_after = Pt(6)
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell_p = table.cell(0, 0).paragraphs[0]
            cell_p.paragraph_format.line_spacing = 1.0
            cell_p.paragraph_format.space_after = Pt(0)
            
            cleaned_body = "\n".join([l.strip() for l in content_lines if l.strip()]).replace("**", "")
            body_run = cell_p.add_run(cleaned_body)
            body_run.font.size = Pt(12)
            
            s_par = doc.add_paragraph()
            s_par.paragraph_format.space_after = Pt(6)

    # 4. HOD Approval Page
    doc.add_page_break()
    p_hod = doc.add_paragraph()
    run_hod = p_hod.add_run("HOD APPROVAL & REMARKS")
    run_hod.bold = True
    run_hod.font.size = Pt(14)
    p_hod.paragraph_format.space_after = Pt(4)
    
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    
    headers_hod = [("REMARK", 0, 0), ("SIGNATURE / STAMP", 0, 1), ("DATE:", 2, 0), ("NAME:", 2, 1)]
    for text_val, r_i, c_i in headers_hod:
        hp = hod_table.cell(r_i, c_i).paragraphs[0]
        hp.paragraph_format.line_spacing = 1.0
        hp.paragraph_format.space_after = Pt(0)
        hrun = hp.add_run(text_val)
        hrun.bold = True
        hrun.font.size = Pt(12)
        
    hod_table.rows[1].height = Pt(40)

    # Apply tight layout format structure configurations to all table elements
    for t in [admin_table, hod_table]:
        for row in t.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 4. MAIN GUI INTERFACE ---
st.write("---")
u_topic = st.text_input("LESSON TOPIC:")
u_syllabus = st.text_input("SYLLABUS CODE:")
u_extra = st.text_area("SPECIFIC CONTEXT / KEYWORDS (OPTIONAL):")

if st.button("🚀 GENERATE PEDATI LESSON PLAN", type="primary"):
    if not user_api_key:
        st.error("❌ KEY CONFIGURATION ERROR! PLEASE INPUT YOUR GOOGLE GEMINI API KEY AT THE TOP OF THE PAGE FIRST.")
    elif not u_topic or not u_syllabus:
        st.error("❌ PLEASE PROVIDE BOTH A LESSON TOPIC AND A SYLLABUS CODE.")
    else:
        with st.spinner("AI IS BUILDING YOUR PEDATI PLAN..."):
            result = generate_pedati_plan(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name)
            st.session_state['pedati_out'] = result

if 'pedati_out' in st.session_state:
    st.divider()
    st.subheader("👁️ AI PREVIEW")
    st.text_area("GENERATED CONTENT CONTENT PREVIEW", st.session_state['pedati_out'], height=350)
    
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['pedati_out'])
    st.download_button(
        label="📥 DOWNLOAD WORD (.DOCX)", 
        data=doc_file, 
        file_name=f"PEDATI_LP_{u_topic.upper().replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- FOOTER SECTION ---
st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>SMART PEDATI LESSON PLAN AI-GENERATOR V1.0</b></p>
        <p>CONCEPTUALIZED BY: <b>HAJAH NURUL HAZIQAH @ HJH HARTINI HJ NORDIN</b></p>
        <p>© 2026 BSC H.M IN COMPUTER SCIENCE, UNIVERSITY OF STRATHCLYDE</p>
    </div>
    """,
    unsafe_allow_html=True
)
